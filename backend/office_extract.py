"""Office document text extraction (DOCX/XLSX).

Designed for local use and best-effort ingestion into the KB.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional


def extract_office_text(path: Path, *, max_chars: int = 2_000_000, max_cells: int = 50_000) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext == "docx":
        return _extract_docx(path, max_chars=max_chars)
    if ext == "xlsx":
        return _extract_xlsx(path, max_chars=max_chars, max_cells=max_cells)
    raise ValueError(f"Unsupported office extension: .{ext}")


def _extract_docx(path: Path, *, max_chars: int) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Missing dependency: python-docx") from e

    doc = Document(str(path))
    parts: list[str] = []

    for p in doc.paragraphs or []:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
        if sum(len(x) for x in parts) >= max_chars:
            break

    # Tables (best-effort)
    if sum(len(x) for x in parts) < max_chars:
        for ti, table in enumerate(doc.tables or [], start=1):
            parts.append(f"\n[Table {ti}]")
            for row in table.rows or []:
                cells = []
                for cell in row.cells or []:
                    cells.append(_norm_cell_text(getattr(cell, "text", "")))
                line = "\t".join(cells).strip()
                if line:
                    parts.append(line)
                if sum(len(x) for x in parts) >= max_chars:
                    break
            if sum(len(x) for x in parts) >= max_chars:
                break

    text = "\n".join(parts).strip()
    return text[:max_chars]


def _extract_xlsx(path: Path, *, max_chars: int, max_cells: int) -> str:
    try:
        import openpyxl  # type: ignore[import-not-found]
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Missing dependency: openpyxl") from e

    wb = openpyxl.load_workbook(filename=str(path), read_only=True, data_only=True)
    parts: list[str] = []
    used_cells = 0

    for ws in wb.worksheets:
        parts.append(f"\n[Sheet] {ws.title}")
        # Iterate rows; openpyxl read_only yields tuples of cells.
        for row in ws.iter_rows(values_only=True):
            if used_cells >= max_cells:
                parts.append("[... truncated: max_cells reached ...]")
                break
            values = []
            for v in row:
                used_cells += 1
                values.append(_format_xlsx_value(v))
                if used_cells >= max_cells:
                    break
            line = "\t".join(values).rstrip()
            if line.strip():
                parts.append(line)
            if sum(len(x) for x in parts) >= max_chars:
                parts.append("[... truncated: max_chars reached ...]")
                break
        if used_cells >= max_cells or sum(len(x) for x in parts) >= max_chars:
            break

    try:
        wb.close()
    except Exception:
        pass

    text = "\n".join(parts).strip()
    return text[:max_chars]


def _format_xlsx_value(v) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "TRUE" if v else "FALSE"
    s = str(v)
    return _norm_cell_text(s)


def _norm_cell_text(s: Optional[str]) -> str:
    t = str(s or "")
    # Collapse newlines/tabs to keep TSV stable.
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = t.replace("\t", " ").replace("\n", " ").strip()
    return t

